import tempfile
from pathlib import Path
from unittest.mock import patch
from decimal import Decimal
from datetime import date
import uuid

from docx import Document

from app.services.contract_doc_service import render_standard_contract_draft
from app.models.contract import Contract, ContractTemplate
from app.models.customer import Customer
from app.models.service_type import ServiceType


def test_render_standard_contract_draft(db_session):
    # Seed service type with unique code
    unique_code = f"eval-{uuid.uuid4().hex[:8]}"
    st = ServiceType(code=unique_code, name="安全评价", is_active=True)
    db_session.add(st)
    db_session.flush()

    unique_suffix = uuid.uuid4().hex[:8]

    # Seed customer
    customer = Customer(name=f"测试甲方-{unique_suffix}", city="北京市", district="朝阳区", address="测试路1号")
    db_session.add(customer)
    db_session.flush()

    # Seed contract
    contract = Contract(
        contract_no=f"HT-{unique_suffix}",
        title="测试合同",
        customer_id=customer.id,
        service_type=st.id,
        total_amount=Decimal("8888.88"),
        sign_date=date(2024, 6, 1),
        start_date=date(2024, 6, 1),
        end_date=date(2025, 5, 31),
    )
    db_session.add(contract)
    db_session.flush()

    # Seed default template
    template = ContractTemplate(
        name=f"默认模板-{unique_suffix}",
        service_type=st.id,
        file_url="contract-templates/default_template.docx",
        is_default=True,
    )
    db_session.add(template)
    db_session.commit()

    # Build a minimal .docx with placeholders
    with tempfile.TemporaryDirectory() as tmpdir:
        tpl_path = Path(tmpdir) / "tpl.docx"
        doc = Document()
        doc.add_paragraph("甲方: {{party_a_name}}")
        doc.add_paragraph("金额大写: {{total_amount_upper}}")
        doc.save(str(tpl_path))
        template_bytes = tpl_path.read_bytes()

    with patch("app.services.contract_doc_service._download_minio_file", return_value=template_bytes):
        with patch("app.services.contract_doc_service._upload_bytes_to_minio") as mock_upload:
            result = render_standard_contract_draft(contract, db_session)
            assert result is not None
            assert result.startswith(f"contracts/{contract.id}/standard_drafts/")
            assert result.endswith(".docx")
            assert mock_upload.called
            uploaded_bytes = mock_upload.call_args[0][0]
            assert isinstance(uploaded_bytes, bytes)
            assert len(uploaded_bytes) > 0
