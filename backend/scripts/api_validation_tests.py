#!/usr/bin/env python3
"""
API 验证测试脚本 — 验证批次 1/2 修复点
"""

import concurrent.futures
import uuid
from datetime import date, timedelta

import requests

BASE_URL = "http://localhost:8000/api/v1"
ADMIN_USER = {"username": "admin", "password": "Admin@123456"}

session = requests.Session()


def login():
    r = session.post(f"{BASE_URL}/auth/login", json=ADMIN_USER)
    assert r.status_code == 200, f"登录失败: {r.text}"
    data = r.json()
    session.headers["Authorization"] = f"Bearer {data['access_token']}"
    # 验证 cookie 中存在 token
    cookies = session.cookies.get_dict()
    assert "access_token" in cookies, "Cookie 中未下发 access_token"
    print("✅ 登录成功，Cookie 已下发 token")
    return data


def test_cookie_auth():
    """验证 Cookie 认证方式"""
    s = requests.Session()
    # 先登录让 cookie 生效
    r = s.post(f"{BASE_URL}/auth/login", json=ADMIN_USER)
    assert r.status_code == 200
    # 清除 header，仅依赖 cookie
    s.headers.pop("Authorization", None)
    r = s.get(f"{BASE_URL}/auth/me")
    assert r.status_code == 200, f"Cookie 认证失败: {r.text}"
    print("✅ Cookie 认证方式正常工作")


def create_customer() -> int:
    uid = str(uuid.uuid4())[:8]
    payload = {
        "name": f"测试客户-{uid}",
        "credit_code": f"91110105MA{uid}",
        "industry": "IT",
        "address": "北京市",
        "contacts": [{"name": "张三", "phone": "13800138000", "email": "zhangsan@example.com"}],
    }
    r = session.post(f"{BASE_URL}/customers", json=payload)
    assert r.status_code == 201, f"创建客户失败: {r.text}"
    return r.json()["id"]


SERVICE_TYPE_ID = None


def _db_conn():
    import os

    import psycopg2

    return psycopg2.connect(
        host=os.getenv("DB_HOST", "localhost"),
        port=os.getenv("DB_PORT", "5432"),
        dbname=os.getenv("DB_NAME", "safety_bms"),
        user=os.getenv("DB_USER", "postgres"),
        password=os.getenv("DB_PASSWORD", "postgres"),
    )


def get_service_type_id() -> int:
    global SERVICE_TYPE_ID
    if SERVICE_TYPE_ID is None:
        conn = _db_conn()
        cur = conn.cursor()
        cur.execute("SELECT id FROM service_types WHERE is_active = TRUE LIMIT 1")
        row = cur.fetchone()
        if row:
            SERVICE_TYPE_ID = row[0]
        else:
            cur.execute(
                "INSERT INTO service_types (code, name, is_active, created_at, updated_at) VALUES (%s, %s, TRUE, NOW(), NOW()) RETURNING id",
                ("evaluation", "安全评价"),
            )
            row = cur.fetchone()
            assert row is not None
            SERVICE_TYPE_ID = row[0]
            conn.commit()
        cur.close()
        conn.close()
    assert SERVICE_TYPE_ID is not None
    return SERVICE_TYPE_ID


def ensure_test_template() -> int:
    """确保存在一个可用的合同模板，没有则自动创建并上传一个最小 docx 文件。"""
    r = session.get(f"{BASE_URL}/contract-templates?page=1&page_size=1")
    if r.status_code == 200 and r.json().get("items"):
        return r.json()["items"][0]["id"]

    # 创建一个最小 docx 文件
    from docx import Document

    doc = Document()
    doc.add_paragraph("合同编号：{{contract_no}}")
    doc.add_paragraph("客户名称：{{customer_name}}")
    doc_path = "/tmp/test_contract_template.docx"
    doc.save(doc_path)

    service_type_id = get_service_type_id()
    conn = _db_conn()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO contract_templates (name, service_type, file_url, is_default, created_by, created_at, updated_at) VALUES (%s, %s, %s, %s, %s, NOW(), NOW()) RETURNING id",
        ("测试模板", service_type_id, "contract-templates/test.docx", False, 1),
    )
    row = cur.fetchone()
    assert row is not None
    template_id = row[0]
    conn.commit()
    cur.close()
    conn.close()

    # 上传模板文件
    with open(doc_path, "rb") as f:
        upload_r = session.post(
            f"{BASE_URL}/contract-templates/{template_id}/upload",
            files={
                "file": (
                    "test_contract_template.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert upload_r.status_code == 200, f"上传测试模板文件失败: {upload_r.text}"
    return template_id


def create_contract(customer_id: int, total_amount: float = 10000, template_id: int = None):
    payload = {
        "contract_no": f"C-{date.today().isoformat()}-{total_amount}-{str(uuid.uuid4())[:4]}",
        "title": "测试合同",
        "customer_id": customer_id,
        "service_type": get_service_type_id(),
        "total_amount": total_amount,
        "sign_date": date.today().isoformat(),
        "start_date": date.today().isoformat(),
        "end_date": (date.today() + timedelta(days=30)).isoformat(),
        "status": "active",
    }
    if template_id:
        payload["template_id"] = template_id
    r = session.post(f"{BASE_URL}/contracts", json=payload)
    assert r.status_code == 201, f"创建合同失败: {r.text}"
    return r.json()["id"]


def test_invoice_amount_limit(contract_id: int):
    """开票金额不能超过合同总额，且并发下不会超额"""
    payload = {
        "contract_id": contract_id,
        "invoice_no": f"INV-{contract_id}-001",
        "invoice_type": "special",
        "amount": 6000,
        "tax_rate": 0.06,
        "invoice_date": date.today().isoformat(),
    }
    r1 = session.post(f"{BASE_URL}/invoices", json=payload)
    assert r1.status_code == 201, f"首次开票失败: {r1.text}"

    # 第二次开 6000，应失败（总额 10000）
    payload["invoice_no"] = f"INV-{contract_id}-002"
    r2 = session.post(f"{BASE_URL}/invoices", json=payload)
    assert r2.status_code == 400, f"超额开票应被拒绝: {r2.text}"
    print("✅ 开票金额上限校验通过")

    # 并发测试：再开两个 4000，只有一个应成功
    def inv(i):
        p = payload.copy()
        p["invoice_no"] = f"INV-{contract_id}-CC-{i}"
        p["amount"] = 4000
        return session.post(f"{BASE_URL}/invoices", json=p)

    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as ex:
        results = list(ex.map(inv, range(2)))
    ok = sum(1 for x in results if x.status_code == 201)
    assert ok <= 1, f"并发开票竞态: 成功 {ok} 次"
    print("✅ 开票并发竞态条件修复通过")


def test_payment_amount_limit(contract_id: int):
    """收款金额不能超过合同总额，且并发下不会超额"""
    payload = {
        "contract_id": contract_id,
        "payment_no": f"PAY-{contract_id}-001",
        "amount": 7000,
        "payment_date": date.today().isoformat(),
        "payment_method": "bank_transfer",
    }
    r1 = session.post(f"{BASE_URL}/payments", json=payload)
    assert r1.status_code == 201, f"首次收款失败: {r1.text}"

    payload["payment_no"] = f"PAY-{contract_id}-002"
    r2 = session.post(f"{BASE_URL}/payments", json=payload)
    assert r2.status_code == 400, f"超额收款应被拒绝: {r2.text}"
    print("✅ 收款金额上限校验通过")

    def pay(i):
        p = payload.copy()
        p["payment_no"] = f"PAY-{contract_id}-CC-{i}"
        p["amount"] = 4000
        return session.post(f"{BASE_URL}/payments", json=p)

    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as ex:
        results = list(ex.map(pay, range(2)))
    ok = sum(1 for x in results if x.status_code == 201)
    assert ok <= 1, f"并发收款竞态: 成功 {ok} 次"
    print("✅ 收款并发竞态条件修复通过")


def activate_contract(contract_id: int):
    # draft -> review -> active
    r = session.post(f"{BASE_URL}/contracts/{contract_id}/status", json={"status": "review"})
    assert r.status_code == 200, f"合同转审核失败: {r.text}"
    r = session.post(f"{BASE_URL}/contracts/{contract_id}/status", json={"status": "active"})
    assert r.status_code == 200, f"合同激活失败: {r.text}"


def test_contract_status_machine(contract_id: int):
    """合同非法状态转移被阻止"""
    activate_contract(contract_id)
    # active -> draft 应为非法
    r = session.post(f"{BASE_URL}/contracts/{contract_id}/status", json={"status": "draft"})
    assert r.status_code == 400, f"非法状态转移应被拒绝: {r.text}"
    print("✅ 合同状态机校验通过")


def test_contract_audit_to_sign_flow(customer_id: int):
    """合同审核-签订融合流程：提交审核自动生草稿，active 状态不可修改，支持上传盖章版签订"""
    # 1. 查询一个可用模板
    r = session.get(f"{BASE_URL}/contract-templates?page=1&page_size=1")
    template_id = None
    if r.status_code == 200 and r.json().get("items"):
        template_id = r.json()["items"][0]["id"]

    contract_id = create_contract(customer_id, total_amount=5000)

    # 2. 不带模板提交审核应失败
    r = session.post(f"{BASE_URL}/contracts/{contract_id}/status", json={"status": "review"})
    if template_id is None:
        print("⚠️  无可用模板，跳过部分测试")
        # 直接激活并走原有流程
        r = session.post(f"{BASE_URL}/contracts/{contract_id}/status", json={"status": "review"})
        assert r.status_code == 200
        r = session.post(f"{BASE_URL}/contracts/{contract_id}/status", json={"status": "active"})
        assert r.status_code == 200
        return contract_id

    assert r.status_code == 400, f"无模板提交审核应被拒绝: {r.text}"
    print("✅ 无模板提交审核被拒绝")

    # 3. 关联模板后提交审核，应自动生成草稿
    r = session.patch(f"{BASE_URL}/contracts/{contract_id}", json={"template_id": template_id})
    assert r.status_code == 200, f"关联模板失败: {r.text}"

    r = session.post(f"{BASE_URL}/contracts/{contract_id}/status", json={"status": "review"})
    assert r.status_code == 200, f"提交审核失败: {r.text}"
    data = r.json()
    assert data["draft_doc_url"] is not None, "提交审核后未生成草稿"
    print("✅ 提交审核自动生成草稿通过")

    # 4. review 状态不可修改
    r = session.patch(f"{BASE_URL}/contracts/{contract_id}", json={"total_amount": 1})
    assert r.status_code == 400, f"review 状态应禁止修改: {r.text}"
    print("✅ review 状态锁定内容通过")

    # 5. 审核通过 -> active
    r = session.post(f"{BASE_URL}/contracts/{contract_id}/status", json={"status": "active"})
    assert r.status_code == 200, f"审核通过失败: {r.text}"

    # 6. active 状态不可修改
    r = session.patch(f"{BASE_URL}/contracts/{contract_id}", json={"total_amount": 1})
    assert r.status_code == 400, f"active 状态应禁止修改: {r.text}"
    print("✅ active 状态锁定内容通过")

    # 7. 上传盖章版签订（模拟文件路径）
    r = session.post(
        f"{BASE_URL}/contracts/{contract_id}/upload-signed",
        json={"file_url": f"contracts/{contract_id}/finals/signed_test.pdf"},
    )
    assert r.status_code == 200, f"上传盖章版签订失败: {r.text}"
    assert r.json()["status"] == "signed"
    assert r.json()["final_pdf_url"] is not None
    print("✅ 上传盖章版签订流程通过")

    return contract_id


def test_admin_self_lockout(user_id: int):
    """Admin 不能禁用/降权自己"""
    r = session.patch(f"{BASE_URL}/users/{user_id}", json={"is_active": False})
    assert r.status_code == 400, f"Admin 自禁用应被拒绝: {r.text}"
    print("✅ Admin 自禁用防护通过")


def test_weak_password_rejected():
    """弱密码创建用户被拒绝"""
    r = session.post(
        f"{BASE_URL}/users",
        json={"username": "weakuser", "password": "123", "email": "weak@example.com"},
    )
    assert r.status_code == 422, f"弱密码应被拒绝: {r.text}"
    print("✅ 弱密码校验通过")


def test_minio_invalid_file(contract_id: int):
    """MinIO 非法文件上传被拒绝"""
    r = session.post(
        f"{BASE_URL}/contracts/{contract_id}/upload",
        files={"file": ("test.exe", b"fake exe content", "application/octet-stream")},
    )
    assert r.status_code == 400, f"非法扩展名应被拒绝: {r.text}"

    # 超大文件（>10MB）
    big = b"x" * (10 * 1024 * 1024 + 1)
    r2 = session.post(
        f"{BASE_URL}/contracts/{contract_id}/upload",
        files={"file": ("test.pdf", big, "application/pdf")},
    )
    assert r2.status_code == 400, f"超大文件应被拒绝: {r2.text}"
    print("✅ MinIO 文件上传校验通过")


def get_me():
    r = session.get(f"{BASE_URL}/auth/me")
    assert r.status_code == 200
    return r.json()


def main():
    print("=" * 50)
    print("开始 API 验证测试")
    print("=" * 50)

    login()
    test_cookie_auth()

    me = get_me()
    test_admin_self_lockout(me["id"])
    test_weak_password_rejected()

    customer_id = create_customer()
    template_id = ensure_test_template()

    contract_id = create_contract(customer_id, total_amount=10000, template_id=template_id)
    test_contract_status_machine(contract_id)  # 内部会 activate_contract
    test_invoice_amount_limit(contract_id)

    contract_id2 = create_contract(customer_id, total_amount=10000, template_id=template_id)
    activate_contract(contract_id2)
    test_payment_amount_limit(contract_id2)

    test_contract_audit_to_sign_flow(customer_id)

    contract_id3 = create_contract(customer_id, total_amount=100, template_id=template_id)
    activate_contract(contract_id3)

    print("=" * 50)
    print("🎉 全部 API 验证测试通过")
    print("=" * 50)


if __name__ == "__main__":
    main()
