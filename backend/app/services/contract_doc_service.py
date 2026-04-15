import subprocess
import tempfile
import uuid
from pathlib import Path

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches

import logging

from app.models.contract import ContractTemplate
from app.services.minio_service import minio_service


def _download_minio_file(object_name: str) -> bytes:
    """从 MinIO 下载文件并返回 bytes"""
    response = minio_service.client.get_object(minio_service.bucket, object_name)
    return response.read()


def _upload_bytes_to_minio(data: bytes, object_name: str, content_type: str = "application/octet-stream") -> str:
    """上传 bytes 到 MinIO，返回 object_name"""
    from io import BytesIO

    minio_service.client.put_object(
        minio_service.bucket,
        object_name,
        BytesIO(data),
        length=len(data),
        content_type=content_type,
    )
    return object_name


def render_contract_draft(contract, template_object_name: str) -> str:
    """
    根据合同数据和 .docx 模板生成待签合同文档，上传到 MinIO，返回 object_name
    """
    # 下载模板
    template_bytes = _download_minio_file(template_object_name)

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "template.docx"
        template_path.write_bytes(template_bytes)

        doc = DocxTemplate(str(template_path))

        context = {
            "contract_no": contract.contract_no or "",
            "title": contract.title or "",
            "customer_name": contract.customer.name if contract.customer else "",
            "service_type_label": _get_service_type_label(contract),
            "total_amount": str(contract.total_amount) if contract.total_amount else "0.00",
            "payment_plan_label": _get_payment_plan_label(contract.payment_plan.value if contract.payment_plan else "once"),
            "start_date": contract.start_date.isoformat() if contract.start_date else "",
            "end_date": contract.end_date.isoformat() if contract.end_date else "",
            "sign_date": contract.sign_date.isoformat() if contract.sign_date else "",
            "remark": contract.remark or "",
            "created_at": contract.created_at.isoformat() if contract.created_at else "",
        }

        doc.render(context)

        output_path = Path(tmpdir) / "draft.docx"
        doc.save(str(output_path))

        draft_bytes = output_path.read_bytes()

    # 上传到 MinIO
    draft_object_name = f"contracts/{contract.id}/drafts/{uuid.uuid4().hex}.docx"
    _upload_bytes_to_minio(draft_bytes, draft_object_name, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return draft_object_name


def insert_signatures_and_to_pdf(
    contract,
    party_a_name: str,
    party_a_signature_object_name: str,
    party_b_name: str,
    party_b_signature_object_name: str,
) -> str:
    """
    将签名图片插入待签文档，转换为 PDF，上传到 MinIO，返回最终 PDF 的 object_name
    """
    if not contract.draft_doc_url:
        raise ValueError("合同没有待签文档")

    draft_bytes = _download_minio_file(contract.draft_doc_url)
    sig_a_bytes = _download_minio_file(party_a_signature_object_name)
    sig_b_bytes = _download_minio_file(party_b_signature_object_name)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        draft_path = tmpdir_path / "draft.docx"
        draft_path.write_bytes(draft_bytes)

        sig_a_path = tmpdir_path / "sig_a.png"
        sig_a_path.write_bytes(sig_a_bytes)

        sig_b_path = tmpdir_path / "sig_b.png"
        sig_b_path.write_bytes(sig_b_bytes)

        # 打开文档，插入签名
        doc = Document(str(draft_path))
        _insert_signature_placeholders(doc, party_a_name, str(sig_a_path), party_b_name, str(sig_b_path))

        signed_docx_path = tmpdir_path / "signed.docx"
        doc.save(str(signed_docx_path))

        # LibreOffice 转 PDF
        pdf_path = _convert_docx_to_pdf(str(signed_docx_path), str(tmpdir_path))
        pdf_bytes = Path(pdf_path).read_bytes()

    pdf_object_name = f"contracts/{contract.id}/finals/{uuid.uuid4().hex}.pdf"
    _upload_bytes_to_minio(pdf_bytes, pdf_object_name, content_type="application/pdf")
    return pdf_object_name


def _insert_signature_placeholders(doc: Document, party_a_name: str, sig_a_path: str, party_b_name: str, sig_b_path: str):
    """
    在文档末尾添加签名区域。如果文档中已有特定占位文本，则替换为签名图片。
    """
    # 先尝试替换正文中的占位符
    replaced_a = _replace_placeholder_with_image(doc, "{{signature_party_a}}", sig_a_path)
    replaced_b = _replace_placeholder_with_image(doc, "{{signature_party_b}}", sig_b_path)

    if not replaced_a or not replaced_b:
        # 如果没有找到占位符，在文档末尾追加签名块
        doc.add_paragraph()
        doc.add_paragraph("甲方（签章）：")
        if not replaced_a:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(sig_a_path, width=Inches(1.5))
            doc.add_paragraph(f"签署人：{party_a_name}")
        doc.add_paragraph()
        doc.add_paragraph("乙方（签章）：")
        if not replaced_b:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(sig_b_path, width=Inches(1.5))
            doc.add_paragraph(f"签署人：{party_b_name}")


def _replace_placeholder_with_image(doc: Document, placeholder: str, image_path: str) -> bool:
    """
    遍历文档段落，如果某段仅包含占位文本，则清空并插入图片。
    """
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(1.5))
            return True
    # 也检查表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.clear()
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(1.5))
                        return True
    return False


def _convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
    """
    调用 LibreOffice headless 将 docx 转换为 pdf
    """
    cmd = [
        "soffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        output_dir,
        docx_path,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice 转换 PDF 失败: {result.stderr}")

    base_name = Path(docx_path).stem
    pdf_path = Path(output_dir) / f"{base_name}.pdf"
    if not pdf_path.exists():
        raise RuntimeError(f"PDF 文件未生成: {pdf_path}")
    return str(pdf_path)


def save_base64_signature_to_minio(base64_data: str, prefix: str, max_size_mb: float = 2.0) -> str:
    """
    将 base64 签名图片保存到 MinIO，返回 object_name
    """
    import base64
    from io import BytesIO

    # 去掉可能存在的 data:image/png;base64, 前缀
    if "," in base64_data:
        base64_data = base64_data.split(",", 1)[1]

    image_bytes = base64.b64decode(base64_data)
    max_size = int(max_size_mb * 1024 * 1024)
    if len(image_bytes) > max_size:
        raise ValueError(f"签名图片大小超过 {max_size_mb}MB 限制")

    object_name = f"{prefix}/{uuid.uuid4().hex}.png"

    minio_service.client.put_object(
        minio_service.bucket,
        object_name,
        BytesIO(image_bytes),
        length=len(image_bytes),
        content_type="image/png",
    )
    return object_name


def _get_service_type_label(contract) -> str:
    if contract.service_type_obj:
        return contract.service_type_obj.name
    return ""


def _get_payment_plan_label(value: str) -> str:
    labels = {
        "once": "一次性",
        "installment": "分期",
    }
    return labels.get(value, value)


def number_to_chinese_upper(amount) -> str:
    from decimal import Decimal, ROUND_HALF_UP

    amount = Decimal(str(amount)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    integer_part = int(amount)
    decimal_part = int((amount - integer_part) * 100)

    nums = "零壹贰叁肆伍陆柒捌玖"
    units = ["", "拾", "佰", "仟"]

    def _four_digit_to_chinese(n: int) -> str:
        if n == 0:
            return "零"
        s = str(n)
        result = []
        for i, ch in enumerate(s):
            digit = int(ch)
            pos = len(s) - i - 1
            if digit == 0:
                if result and result[-1] != "零":
                    result.append("零")
            else:
                result.append(nums[digit] + units[pos % 4])
        return "".join(result).rstrip("零")

    def _int_to_chinese(n: int) -> str:
        if n == 0:
            return "零"
        if n < 10000:
            return _four_digit_to_chinese(n)

        low = n % 10000
        mid = (n // 10000) % 10000
        high = n // 100000000

        parts = []
        if high > 0:
            parts.append(_four_digit_to_chinese(high) + "亿")
        if mid > 0:
            parts.append(_four_digit_to_chinese(mid) + "万")
        elif high > 0 and low > 0:
            parts.append("零")

        if low > 0:
            low_str = _four_digit_to_chinese(low)
            if mid > 0 and len(str(low)) < 4 and not low_str.startswith("零"):
                low_str = "零" + low_str
            parts.append(low_str)

        return "".join(parts).rstrip("零")

    jiao = decimal_part // 10
    fen = decimal_part % 10

    if integer_part == 0:
        if jiao == 0 and fen == 0:
            return "零元整"
        result = ""
    else:
        result = _int_to_chinese(integer_part) + "元"

    if jiao == 0 and fen == 0:
        result += "整"
    else:
        if jiao > 0:
            result += nums[jiao] + "角"
        if fen > 0:
            if integer_part > 0 and jiao == 0:
                result += "零"
            result += nums[fen] + "分"

    return result


def _build_standard_contract_context(contract) -> dict:
    customer = contract.customer
    sign_date = contract.sign_date
    start_date = contract.start_date
    end_date = contract.end_date

    sign_location = ""
    if customer:
        sign_location = f"{customer.city or ''}{customer.district or ''}"

    total_amount = contract.total_amount or 0

    return {
        "contract_reg_no": contract.contract_no or "",
        "party_a_name": customer.name if customer else "",
        "sign_location": sign_location,
        "sign_year": sign_date.year if sign_date else "",
        "sign_month": sign_date.month if sign_date else "",
        "sign_day": sign_date.day if sign_date else "",
        "valid_start_year": start_date.year if start_date else "",
        "valid_start_month": start_date.month if start_date else "",
        "valid_start_day": start_date.day if start_date else "",
        "valid_end_year": end_date.year if end_date else "",
        "valid_end_month": end_date.month if end_date else "",
        "valid_end_day": end_date.day if end_date else "",
        "service_start_year": start_date.year if start_date else "",
        "service_start_month": start_date.month if start_date else "",
        "service_start_day": start_date.day if start_date else "",
        "service_end_year": end_date.year if end_date else "",
        "service_end_month": end_date.month if end_date else "",
        "service_end_day": end_date.day if end_date else "",
        "total_amount": str(total_amount),
        "total_amount_upper": number_to_chinese_upper(total_amount),
        "payment_amount": str(total_amount),
        "service_address": customer.address if customer else "",
    }


def render_standard_contract_draft(contract, db) -> str | None:
    logger = logging.getLogger(__name__)

    template = db.query(ContractTemplate).filter(ContractTemplate.is_default == True).first()
    if not template or not template.file_url:
        logger.warning("未找到默认合同模板，跳过标准合同草稿生成")
        return None

    try:
        template_bytes = _download_minio_file(template.file_url)

        with tempfile.TemporaryDirectory() as tmpdir:
            template_path = Path(tmpdir) / "template.docx"
            template_path.write_bytes(template_bytes)

            doc = DocxTemplate(str(template_path))
            context = _build_standard_contract_context(contract)
            doc.render(context)

            output_path = Path(tmpdir) / "standard_draft.docx"
            doc.save(str(output_path))

            draft_bytes = output_path.read_bytes()

        draft_object_name = f"contracts/{contract.id}/standard_drafts/{uuid.uuid4().hex}.docx"
        _upload_bytes_to_minio(
            draft_bytes,
            draft_object_name,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        return draft_object_name
    except Exception:
        logger.exception("标准合同草稿生成失败")
        return None
